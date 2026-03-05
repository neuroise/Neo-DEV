#!/usr/bin/env python3
"""
Generate D1 — Documento di Analisi (Word .docx)

Contract deliverable: "Report sull'architettura esistente con raccomandazioni di miglioramento"
Output: deliverables/D1_Documento_Analisi.docx
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

BASE_DIR = Path(__file__).parent.parent
OUT_DIR = BASE_DIR / "deliverables"
OUT_DIR.mkdir(exist_ok=True)

# ── Styling helpers ──────────────────────────────────────────────────────────

def set_style(doc):
    """Configure base document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            h.font.size = Pt(18)
        elif level == 2:
            h.font.size = Pt(14)
        else:
            h.font.size = Pt(12)


def add_cover(doc):
    """Add a cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NEURØISE")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("D1 — Documento di Analisi")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Report sull'architettura esistente con raccomandazioni di miglioramento")
    run.font.size = Pt(12)
    run.italic = True

    for _ in range(4):
        doc.add_paragraph()

    info = [
        ("Progetto", "NEURØISE — Intelligent Storytelling Engine for Luxury Experiences"),
        ("Contratto", "Collaborazione di Ricerca — Università di Pisa / NO NOISE S.r.l."),
        ("Autore", "Prof. Tiberio Uricchio — Università di Pisa"),
        ("Data", "Marzo 2026"),
        ("Versione", "1.0"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(value)
        run.font.size = Pt(10)

    doc.add_page_break()


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


# ── Content ──────────────────────────────────────────────────────────────────

def section_1_introduzione(doc):
    doc.add_heading("1. Introduzione", level=1)

    doc.add_heading("1.1 Contesto del Progetto", level=2)
    doc.add_paragraph(
        "NEURØISE è un motore di storytelling intelligente progettato per creare "
        "esperienze personalizzate nel settore luxury, con focus iniziale sulle "
        "esperienze a bordo di yacht di lusso. Il sistema genera contenuti video "
        "e musicali su misura basandosi sul profilo psicologico dell'ospite, "
        "trasformando dati comportamentali in narrazioni audiovisive emozionali."
    )
    doc.add_paragraph(
        "Il Framework v2, sviluppato da NO NOISE, definisce l'architettura concettuale "
        "del sistema attraverso il modello \"Cognitive Sandwich\": un approccio a tre "
        "livelli (Input → Decision → Control → Production) che guida la generazione "
        "di contenuti personalizzati partendo dalla profilazione dell'utente fino "
        "alla produzione multimediale."
    )

    doc.add_heading("1.2 Obiettivo della Collaborazione", level=2)
    doc.add_paragraph(
        "La collaborazione tra l'Università di Pisa e NO NOISE S.r.l. ha l'obiettivo "
        "di sviluppare un Minimum Viable Product (MVP) del motore NEURØISE, validando "
        "le scelte architetturali del Framework v2 attraverso un Playground sperimentale "
        "che consenta la generazione, valutazione e iterazione rapida dei contenuti."
    )

    doc.add_heading("1.3 Ambito dell'Analisi", level=2)
    doc.add_paragraph(
        "Il presente documento si concentra sull'analisi dei seguenti componenti chiave:"
    )
    items = [
        "Director Agent — il modulo LLM responsabile della generazione creativa",
        "Policy Gate — il sistema di validazione e controllo qualità",
        "Framework di Metriche — il sistema di valutazione automatica a 13 dimensioni",
        "Pipeline di Generazione Video — il servizio di text-to-video",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def section_2_architettura(doc):
    doc.add_heading("2. Analisi dell'Architettura Esistente", level=1)

    doc.add_heading("2.1 Modello Cognitive Sandwich", level=2)
    doc.add_paragraph(
        "L'architettura NEURØISE si basa sul modello Cognitive Sandwich definito "
        "nel Framework v2. Questo modello organizza il flusso di elaborazione in "
        "quattro livelli gerarchici:"
    )

    layers = [
        ("Input Layer", "Acquisizione dati utente (profilo psicologico, preferenze musicali, "
         "contesto ambientale). Nel Playground, questo è rappresentato dai profili JSON "
         "con archetipo primario, seed musicale e hint narrativo."),
        ("Decision Layer (Director)", "Un agente LLM che interpreta il profilo utente e "
         "genera decisioni creative sotto forma di un video triptych (3 scene narrative) "
         "e un prompt per la colonna sonora originale (OST)."),
        ("Control Layer (PolicyGate)", "Sistema di validazione rule-based che verifica "
         "la conformità dell'output alle policy del brand (contenuti marini, assenza di "
         "elementi urbani/violenti, coerenza archetipale)."),
        ("Production Layer", "Generazione effettiva dei contenuti multimediali "
         "(video via Wan2.2/TurboWan, musica via modelli generativi)."),
    ]
    for name, desc in layers:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("2.2 Flusso Dati nel Playground", level=2)
    doc.add_paragraph(
        "Il flusso dati implementato nel Playground segue la pipeline:"
    )
    steps = [
        "Profile JSON → caricamento profilo utente con archetipo, seed musicale, hint narrativo",
        "Director LLM → generazione del video_triptych (3 scene) e ost_prompt",
        "PolicyGate → validazione con 8 regole (RED/YELLOW/GREEN)",
        "Metriche Automatiche → calcolo di 13 dimensioni di qualità (M_AUTO_01–13)",
        "Video Generation → text-to-video con Wan2.2 o TurboWan (opzionale)",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("2.3 Schema del Profilo Utente", level=2)
    doc.add_paragraph(
        "Ogni profilo utente è definito in formato JSON con la seguente struttura:"
    )

    # Load example profile
    profile_path = BASE_DIR / "data" / "profiles" / "official" / "S-01.json"
    try:
        with open(profile_path) as f:
            profile = json.load(f)
        p = doc.add_paragraph()
        p.style = "No Spacing"
        run = p.add_run(json.dumps(profile, indent=2))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    except FileNotFoundError:
        doc.add_paragraph("[Profilo S-01.json non trovato]")

    doc.add_paragraph()
    doc.add_paragraph(
        "I 30 profili ufficiali sono suddivisi equamente tra tre archetipi: "
        "Sage (S-01..S-10), Rebel (R-01..R-10) e Lover (L-01..L-10). "
        "Ogni profilo include un archetipo primario, parametri musicali di riferimento "
        "(genere, BPM, tag di mood) e un hint narrativo che guida la coerenza tematica."
    )

    doc.add_heading("2.4 Schema dell'Output Director", level=2)
    doc.add_paragraph(
        "Il Director genera un output JSON strutturato contenente:"
    )
    outputs = [
        ("video_triptych", "Array di 3 scene (start → evolve → end), ciascuna con prompt "
         "di produzione (50-500 caratteri), durata suggerita, tag di mood e hint di camera."),
        ("ost_prompt", "Specifica per la colonna sonora: prompt descrittivo, genere, BPM "
         "(obbligatorio, nell'intervallo dell'archetipo), mood e strumenti suggeriti."),
        ("metadata", "Archetipo rilevato, thread narrativo utilizzato, note di coerenza."),
    ]
    for name, desc in outputs:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        p.add_run(desc)


def section_3_componenti(doc):
    doc.add_heading("3. Componenti Implementati nel Playground", level=1)

    # 3.1 Director
    doc.add_heading("3.1 Director LLM", level=2)
    doc.add_paragraph(
        "Il Director è il componente centrale del sistema, implementato in "
        "core/llm/director.py. Agisce come un Creative Director che interpreta "
        "il profilo psicologico dell'utente e genera contenuti narrativi personalizzati."
    )
    doc.add_paragraph("Caratteristiche principali:")
    features = [
        "Supporto per 3 archetipi junghiani: Sage (contemplativo, 60-80 BPM), "
        "Rebel (dinamico, 120-140 BPM), Lover (intimo, 70-90 BPM)",
        "Output strutturato: video triptych di 3 scene con arco narrativo (start → evolve → end)",
        "Prompt di produzione ottimizzati per text-to-video AI (soggetto, inquadratura, "
        "movimento camera, illuminazione)",
        "Vincolo tematico: esclusivamente ambientazioni marine/costiere",
        "Regole di formato: no riferimenti audio, no metafore, no narrazione in prima persona",
        "Generazione OST con BPM obbligatorio nell'intervallo dell'archetipo",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    # 3.2 PolicyGate
    doc.add_heading("3.2 PolicyGate", level=2)
    doc.add_paragraph(
        "Il PolicyGate (core/gating/policy_gate.py) implementa un sistema di validazione "
        "rule-based con tre livelli di severità:"
    )

    flag_desc = [
        ("GREEN", "Contenuto conforme — procedi con la produzione"),
        ("YELLOW", "Avvertimento — revisione manuale consigliata"),
        ("RED", "Contenuto bloccato — non procedere"),
    ]
    for flag, desc in flag_desc:
        p = doc.add_paragraph()
        run = p.add_run(f"{flag}: ")
        run.bold = True
        run.font.color.rgb = {
            "GREEN": RGBColor(0, 128, 0),
            "YELLOW": RGBColor(200, 150, 0),
            "RED": RGBColor(200, 0, 0),
        }[flag]
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph("Le 8 regole di validazione implementate:")

    rules = [
        ("R001", "Blacklist Check", "RED",
         "Rileva termini proibiti (urban, violence, forest, ecc.) con word boundary matching"),
        ("R002", "Warning Terms", "YELLOW",
         "Segnala termini da verificare (storm, person, logo, ecc.)"),
        ("R003", "Marine Vocabulary", "YELLOW",
         "Verifica presenza di almeno 5 termini marini nel contenuto"),
        ("R004", "Structure Validation", "RED",
         "Valida struttura JSON: video_triptych presente, esattamente 3 scene, ost_prompt completo"),
        ("R005", "Scene Sequence", "RED",
         "Verifica sequenza corretta delle scene: start → evolve → end"),
        ("R006", "Archetype Consistency", "YELLOW",
         "Confronta densità keyword per archetipo, segnala se altro archetipo domina (>1.5×)"),
        ("R007", "Prompt Length", "YELLOW",
         "Verifica lunghezza prompt: ≥50 e ≤500 caratteri per scena"),
        ("R008", "BPM Validation", "RED/YELLOW",
         "RED se BPM mancante nell'OST; YELLOW se fuori range archetipale"),
    ]
    add_table(doc,
              ["Regola", "Nome", "Severità", "Descrizione"],
              [(r[0], r[1], r[2], r[3]) for r in rules],
              col_widths=[2, 3.5, 2.5, 9])

    # 3.3 Adapter LLM
    doc.add_heading("3.3 Adapter LLM", level=2)
    doc.add_paragraph(
        "Il sistema adotta un'architettura modulare per l'integrazione con diversi "
        "provider LLM, implementata attraverso il pattern Adapter (core/llm/):"
    )
    adapters = [
        ("OllamaAdapter", "Per modelli locali (LLaMA 3.3:70b, Qwen3:32b) — "
         "gestione num_ctx per VRAM, temperature clamping, timeout configurabile"),
        ("AnthropicAdapter", "Per Claude (API cloud) — supporto streaming, gestione token"),
        ("OpenAIAdapter", "Per GPT-4/o1 (API cloud) — compatibilità OpenAI standard"),
    ]
    for name, desc in adapters:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        p.add_run(desc)

    doc.add_paragraph(
        "\nLa factory function create_adapter() seleziona automaticamente l'adapter "
        "appropriato in base al nome del modello, consentendo il cambio trasparente "
        "tra provider locali e cloud."
    )

    # 3.4 Prompt Packs
    doc.add_heading("3.4 Prompt Packs", level=2)
    doc.add_paragraph(
        "Per lo studio di ablazione sulla sensibilità del sistema al prompt engineering, "
        "sono stati implementati tre prompt pack (core/llm/prompt_packs/):"
    )
    packs = [
        ("Default", "Prompt completo con istruzioni dettagliate, definizione archetipi, "
         "vincoli di formato e regole di produzione. Rappresenta il baseline."),
        ("Concise", "Versione minimale (~200 parole) con le sole informazioni essenziali: "
         "archetipi sintetici, vincoli core e schema output."),
        ("Detailed", "Versione estesa con tecniche di camera per archetipo, "
         "esempi good/bad, rubrica di autovalutazione a 5 dimensioni."),
    ]
    for name, desc in packs:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    # 3.5 Metriche
    doc.add_heading("3.5 Framework di Metriche Automatiche", level=2)
    doc.add_paragraph(
        "Il framework di valutazione comprende 13 metriche automatiche (M_AUTO_01–13) "
        "organizzate in quattro categorie:"
    )

    metrics = [
        ("Strutturali", [
            ("M_AUTO_01", "Schema Compliance", "Conformità alla struttura JSON richiesta"),
            ("M_AUTO_03", "Role Sequence Valid", "Correttezza sequenza scene (start→evolve→end)"),
            ("M_AUTO_06", "Prompt Length Valid", "Lunghezza prompt nel range 50-500 char"),
            ("M_AUTO_13", "Pacing Progression", "Progressione ritmica nell'arco narrativo"),
        ]),
        ("Semantiche", [
            ("M_AUTO_02", "Archetype Consistency", "Allineamento al profilo archetipale"),
            ("M_AUTO_04", "Story Thread Presence", "Incorporazione dell'hint narrativo"),
            ("M_AUTO_07", "Archetype Lexical Fit", "Densità lessicale per archetipo"),
            ("M_AUTO_10", "Marine Vocabulary Ratio", "Rapporto vocabolario marino/totale"),
        ]),
        ("Qualitative", [
            ("M_AUTO_05", "Red Flag Score", "Assenza di violazioni PolicyGate"),
            ("M_AUTO_09", "Prompt Specificity", "Specificità dei prompt di produzione"),
            ("M_AUTO_08", "Cross-Scene Coherence", "Coerenza semantica tra scene consecutive"),
            ("M_AUTO_11", "Narrative Coherence", "Consistenza narrativa nel triptych"),
        ]),
        ("LLM Judge", [
            ("M_AUTO_12", "LLM Judge Quality",
             "Valutazione multi-dimensionale tramite LLM (5 dimensioni: Visual Clarity, "
             "Archetype Alignment, Narrative Coherence, Emotional Resonance, Marine Adherence)"),
        ]),
    ]

    for cat_name, cat_metrics in metrics:
        doc.add_heading(f"Categoria: {cat_name}", level=3)
        add_table(doc,
                  ["Codice", "Nome", "Descrizione"],
                  [(m[0], m[1], m[2]) for m in cat_metrics],
                  col_widths=[2.5, 4, 10.5])

    doc.add_paragraph(
        "Lo score aggregato è calcolato come media aritmetica di tutte le 13 metriche, "
        "ciascuna normalizzata nell'intervallo [0, 1]."
    )

    # 3.6 Video Generation
    doc.add_heading("3.6 Servizio di Generazione Video", level=2)
    doc.add_paragraph(
        "Il servizio di generazione video (video-gen/) è implementato come microservizio "
        "FastAPI containerizzato con Docker, supportando due pipeline:"
    )
    pipes = [
        ("Wan2.2", "Pipeline standard text-to-video con qualità elevata e tempi di "
         "generazione più lunghi (~2-5 minuti per clip di 5 secondi)"),
        ("TurboWan", "Pipeline ottimizzata per velocità con qualità leggermente ridotta "
         "(~30-60 secondi per clip)"),
    ]
    for name, desc in pipes:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "\nIl servizio è ottimizzato per l'esecuzione su DGX Spark GB10 (Blackwell GPU, "
        "121.7 GB di memoria unificata) e supporta la generazione batch dei triptych video "
        "con configurazione per-scena di risoluzione, frame rate e durata."
    )


def section_4_raccomandazioni(doc):
    doc.add_heading("4. Criticità Identificate e Raccomandazioni", level=1)

    doc.add_paragraph(
        "Dall'analisi dell'architettura e dei risultati sperimentali emergono "
        "le seguenti aree di miglioramento:"
    )

    topics = [
        ("4.1 Profilazione Avanzata",
         "Il sistema attuale utilizza tre archetipi junghiani (Sage/Rebel/Lover) come "
         "base per la personalizzazione. Si raccomanda l'evoluzione verso un modello di "
         "profilazione più sofisticato basato sui Big Five personality traits, integrato "
         "con un agente NeedsProfiler LLM che possa costruire profili dinamici attraverso "
         "interazioni conversazionali con l'utente. Questo approccio consentirebbe una "
         "personalizzazione più granulare rispetto alla categorizzazione discreta in archetipi.",
         "Priorità: Alta — Impatto diretto sulla qualità della personalizzazione"),

        ("4.2 RAG vs Fine-Tuning per Brand Consistency",
         "Per garantire la coerenza con le linee guida del brand, si raccomanda l'adozione "
         "di un sistema RAG (Retrieval-Augmented Generation) che utilizzi un archivio di "
         "esempi approvati e linee guida stilistiche come contesto dinamico per il Director. "
         "Il fine-tuning, sebbene potenzialmente più efficace per la coerenza stilistica, "
         "presenta costi computazionali e di manutenzione significativamente più elevati "
         "e minor flessibilità nel gestire l'evoluzione del brand.",
         "Priorità: Media — Necessario per la produzione, non per l'MVP"),

        ("4.3 Context Window Management per Archivio Narrativo",
         "L'implementazione di un archivio narrativo (\"Narrative Archive\") richiede una "
         "gestione attenta del context window. Con i modelli Ollama configurati a num_ctx=8192 "
         "per ragioni di VRAM, è necessario un sistema di summarizzazione progressiva "
         "(Summarizer Agent) che comprima le esperienze precedenti mantenendo i thread "
         "narrativi rilevanti. Si raccomanda un approccio ibrido: summary compatto nel "
         "context + retrieval dei dettagli completi via RAG.",
         "Priorità: Alta — Fondamentale per la continuità narrativa"),

        ("4.4 Schema Enforcement con Validazione Pydantic",
         "L'attuale parsing dell'output JSON dal Director si basa su regex ed estrazione "
         "manuale. Si raccomanda l'adozione di Pydantic per la validazione rigorosa dello "
         "schema, con modelli tipizzati per video_triptych, ost_prompt e metadata. Questo "
         "ridurrebbe gli errori di parsing (come quelli osservati con Qwen3:32b, dove 2/30 "
         "profili falliscono per errori di tipo) e migliorerebbe la robustezza del sistema.",
         "Priorità: Alta — Riduce fallimenti e migliora affidabilità"),

        ("4.5 Budget e Latency Management",
         "Lo scenario di deployment su yacht richiede attenzione particolare alla gestione "
         "del budget computazionale e della latenza. La connettività limitata in mare aperto "
         "rende necessario il supporto per inferenza edge (modelli locali) con fallback su "
         "API cloud quando disponibile. Si raccomanda un sistema di budget management che "
         "bilanci qualità della generazione e risorse disponibili, con pre-generazione "
         "opportunistica durante periodi di connettività.",
         "Priorità: Media — Critico per il deployment ma non per l'MVP"),

        ("4.6 Persistenza Ibrida",
         "Per il deployment on-yacht, si raccomanda un sistema di persistenza ibrido: "
         "storage locale (SQLite/file-based) per operatività offline con sincronizzazione "
         "automatica verso cloud storage quando disponibile. I profili utente, l'archivio "
         "narrativo e i contenuti generati devono essere disponibili anche in assenza di "
         "connettività, con gestione automatica dei conflitti durante la sincronizzazione.",
         "Priorità: Media — Necessario per il deployment reale"),
    ]

    for title, text, priority in topics:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
        p = doc.add_paragraph()
        run = p.add_run(priority)
        run.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)


def section_5_conclusioni(doc):
    doc.add_heading("5. Conclusioni", level=1)

    doc.add_paragraph(
        "L'analisi dell'architettura NEURØISE rivela un sistema ben strutturato, "
        "basato su principi solidi di separazione delle responsabilità e modularità. "
        "Il modello Cognitive Sandwich fornisce un framework concettuale chiaro per "
        "l'organizzazione della pipeline di generazione."
    )
    doc.add_paragraph(
        "L'implementazione nel Playground ha validato le scelte architetturali "
        "fondamentali, dimostrando la fattibilità dell'approccio con risultati "
        "quantitativi promettenti (score aggregato 0.775 ± 0.049 con LLaMA 3.3:70b "
        "sul benchmark di 30 profili)."
    )
    doc.add_paragraph(
        "Le raccomandazioni presentate nella Sezione 4 identificano le aree chiave "
        "per l'evoluzione del sistema verso un prodotto completo, con priorità "
        "sulla profilazione avanzata, la gestione del context window e il rafforzamento "
        "della validazione dello schema. Queste migliorie sono necessarie per il "
        "passaggio dall'MVP al prodotto di produzione, ma non bloccano la validazione "
        "sperimentale attualmente in corso."
    )
    doc.add_paragraph(
        "Il Playground nella sua forma attuale rappresenta uno strumento efficace per "
        "la sperimentazione e l'iterazione rapida, consentendo la valutazione "
        "sistematica di modelli, prompt e configurazioni attraverso il framework "
        "di metriche automatiche a 13 dimensioni."
    )


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_style(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover(doc)
    section_1_introduzione(doc)
    doc.add_page_break()
    section_2_architettura(doc)
    doc.add_page_break()
    section_3_componenti(doc)
    doc.add_page_break()
    section_4_raccomandazioni(doc)
    doc.add_page_break()
    section_5_conclusioni(doc)

    out_path = OUT_DIR / "D1_Documento_Analisi.docx"
    doc.save(str(out_path))
    print(f"✓ D1 saved to: {out_path}")
    print(f"  Size: {out_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
