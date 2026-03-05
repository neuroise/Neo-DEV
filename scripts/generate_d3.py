#!/usr/bin/env python3
"""
Generate D3 — Relazione Finale (Word .docx) — DRAFT

Contract deliverable: "Documentazione delle attività svolte, risultati ottenuti
e indicazioni per sviluppi futuri"
Output: deliverables/D3_Relazione_Finale_DRAFT.docx
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).parent.parent
OUT_DIR = BASE_DIR / "deliverables"
OUT_DIR.mkdir(exist_ok=True)

# ── Data loading ─────────────────────────────────────────────────────────────

def load_summary(name):
    path = BASE_DIR / "data" / "experiments" / name / "summary.json"
    if path.exists():
        with open(path) as f:
            return json.load(f)
    return None

BASELINE_LLAMA = load_summary("baseline_30_llama70b_v3")
BASELINE_QWEN = load_summary("baseline_30_qwen32b")
ABLATION_CONCISE = load_summary("ablation_concise_llama70b")
ABLATION_DETAILED = load_summary("ablation_detailed_llama70b")

METRIC_NAMES = [
    ("M_AUTO_01_schema_compliance", "Schema Compliance"),
    ("M_AUTO_02_archetype_consistency", "Archetype Consistency"),
    ("M_AUTO_03_role_sequence_valid", "Role Sequence Valid"),
    ("M_AUTO_04_story_thread_presence", "Story Thread Presence"),
    ("M_AUTO_05_red_flag_score", "Red Flag Score"),
    ("M_AUTO_06_prompt_length_valid", "Prompt Length Valid"),
    ("M_AUTO_07_archetype_lexical_fit", "Archetype Lexical Fit"),
    ("M_AUTO_08_cross_scene_coherence", "Cross-Scene Coherence"),
    ("M_AUTO_09_prompt_specificity", "Prompt Specificity"),
    ("M_AUTO_10_marine_vocabulary_ratio", "Marine Vocab Ratio"),
    ("M_AUTO_11_score_narrative_coherence", "Narrative Coherence"),
    ("M_AUTO_12_llm_judge_quality", "LLM Judge Quality"),
    ("M_AUTO_13_pacing_progression", "Pacing Progression"),
    ("aggregate_score", "Aggregate Score"),
]

# ── Styling helpers ──────────────────────────────────────────────────────────

def set_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            h.font.size = Pt(18)
        elif level == 2:
            h.font.size = Pt(14)
        else:
            h.font.size = Pt(12)


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NEURØISE")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("D3 — Relazione Finale")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documentazione delle attività svolte, risultati ottenuti\ne indicazioni per sviluppi futuri")
    run.font.size = Pt(12)
    run.italic = True

    for _ in range(4):
        doc.add_paragraph()
    info = [
        ("Progetto", "NEURØISE — Intelligent Storytelling Engine for Luxury Experiences"),
        ("Contratto", "Collaborazione di Ricerca — Università di Pisa / NO NOISE S.r.l."),
        ("Autore", "Prof. Tiberio Uricchio — Università di Pisa"),
        ("Data", "Marzo 2026"),
        ("Versione", "DRAFT 1.0"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(value)
        run.font.size = Pt(10)
    doc.add_page_break()


def add_table(doc, headers, rows, col_widths=None, bold_last_row=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for r_idx, row_data in enumerate(rows):
        is_last = r_idx == len(rows) - 1
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if bold_last_row and is_last:
                        run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def fmt(val, digits=3):
    """Format a numeric value."""
    if val is None:
        return "—"
    if isinstance(val, str):
        return val
    return f"{val:.{digits}f}"


def get_metric(summary, model_key, metric_key):
    """Extract mean ± std from summary."""
    if summary is None:
        return None, None
    models = summary.get("models", {})
    model_data = models.get(model_key, {})
    m = model_data.get(metric_key, {})
    return m.get("mean"), m.get("std")


def get_archetype_metric(summary, arch, metric_key):
    """Extract mean from archetype breakdown."""
    if summary is None:
        return None
    by_arch = summary.get("by_archetype", {})
    arch_data = by_arch.get(arch, {})
    m = arch_data.get(metric_key, {})
    return m.get("mean")


# ── Sections ─────────────────────────────────────────────────────────────────

def section_1_sommario(doc):
    doc.add_heading("1. Sommario Esecutivo", level=1)
    doc.add_paragraph(
        "Il presente documento riporta le attività svolte, i risultati ottenuti e le "
        "indicazioni per sviluppi futuri nell'ambito della collaborazione di ricerca "
        "tra l'Università di Pisa e NO NOISE S.r.l. per il progetto NEURØISE."
    )
    doc.add_paragraph("Risultati principali:")
    bullets = [
        "Playground sperimentale completamente funzionante con interfaccia Streamlit",
        "30 profili utente ufficiali (10 per archetipo: Sage, Rebel, Lover)",
        "4 esperimenti sistematici completati: baseline, ablation concise, ablation detailed, cross-model",
        "150+ run totali con metriche automatiche a 13 dimensioni",
        "Score aggregato baseline: 0.775 ± 0.049 (LLaMA 3.3:70b, 30/30 successo)",
        "Servizio di generazione video operativo (Wan2.2/TurboWan) su DGX Spark GB10",
        "Framework di valutazione con LLM-as-Judge e predisposizione per valutazione umana",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def section_2_attivita(doc):
    doc.add_heading("2. Attività Svolte", level=1)

    doc.add_heading("2.1 Mese 1 — Analisi e Setup (Dicembre 2025 – Gennaio 2026)", level=2)
    items = [
        "Studio approfondito del Framework v2 (Cognitive Sandwich model)",
        "Analisi dell'architettura esistente e identificazione dei componenti chiave",
        "Setup dell'ambiente di sviluppo: Python, Streamlit, Ollama, Docker",
        "Configurazione DGX Spark GB10: ottimizzazione VRAM (num_ctx=8192), pinning modelli",
        "Progettazione dello schema profilo utente (archetipo, seed musicale, hint narrativo)",
        "Redazione del D1 — Documento di Analisi (v1.0)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.2 Mese 2 — Implementazione Core (Febbraio 2026)", level=2)
    items = [
        "Implementazione Director LLM con supporto per 3 archetipi",
        "Implementazione PolicyGate con 8 regole di validazione (R001-R008)",
        "Architettura adapter modulare: OllamaAdapter, AnthropicAdapter, OpenAIAdapter",
        "Creazione 30 profili ufficiali (S-01..S-10, R-01..R-10, L-01..L-10)",
        "Framework metriche automatiche: 13 dimensioni (M_AUTO_01-13)",
        "Implementazione LLM-as-Judge (qwen3:32b con /no_think)",
        "Interfaccia Streamlit: Home, Generate, Evaluate, Experiments, Analysis, Preview",
        "Experiment runner per batch execution con progress tracking",
        "Rilascio v0.1.0 (23 Febbraio 2026)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.3 Mese 3 — Esperimenti e Deliverables (Marzo 2026)", level=2)
    items = [
        "Baseline experiment: 30 profili × LLaMA 3.3:70b (30/30 successo)",
        "Ablation study: prompt pack concise vs detailed vs default",
        "Cross-model comparison: LLaMA 3.3:70b vs Qwen3:32b",
        "Implementazione prompt packs (default/concise/detailed)",
        "Aggiunta regola R008 (BPM Validation) al PolicyGate",
        "Servizio generazione video: Wan2.2/TurboWan via FastAPI + Docker",
        "Analisi statistica: Wilcoxon signed-rank test, effect size (Cohen's d)",
        "Predisposizione framework valutazione umana (5 dimensioni Likert)",
        "Rilascio v0.2.0 (3 Marzo 2026)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def section_3_architettura(doc):
    doc.add_heading("3. Architettura del Playground", level=1)

    doc.add_paragraph(
        "Il Playground NEURØISE implementa il modello Cognitive Sandwich del Framework v2 "
        "in un'applicazione web interattiva. L'architettura segue il principio di separazione "
        "delle responsabilità con quattro componenti principali:"
    )

    components = [
        ("Director LLM", "Agente creativo che interpreta il profilo utente e genera "
         "video triptych (3 scene) e prompt per colonna sonora. Supporta 3 archetipi: "
         "Sage (contemplativo), Rebel (dinamico), Lover (intimo)."),
        ("PolicyGate", "Sistema di validazione rule-based con 8 regole (R001-R008). "
         "Tre livelli di severità: GREEN (ok), YELLOW (warning), RED (blocco)."),
        ("Framework Metriche", "13 metriche automatiche organizzate in 4 categorie: "
         "strutturali, semantiche, qualitative, LLM Judge."),
        ("Video Generation Service", "Microservizio FastAPI per text-to-video "
         "(Wan2.2/TurboWan) containerizzato con Docker."),
    ]
    for name, desc in components:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("3.1 Stack Tecnologico", level=2)
    stack = [
        ("Linguaggio", "Python 3.12"),
        ("UI", "Streamlit 1.41"),
        ("LLM Runtime", "Ollama (locale), API cloud (Anthropic, OpenAI)"),
        ("Modelli LLM", "LLaMA 3.3:70b (director), Qwen3:32b (judge/comparison)"),
        ("Video Generation", "Wan2.2 / TurboWan via FastAPI"),
        ("Hardware", "NVIDIA DGX Spark GB10 (Blackwell GPU, 121.7 GB unified memory)"),
        ("Containerizzazione", "Docker + Docker Compose"),
        ("Document Generation", "python-docx, python-pptx"),
    ]
    add_table(doc, ["Componente", "Tecnologia"], stack, col_widths=[4, 13])


def section_4_risultati(doc):
    doc.add_heading("4. Risultati Sperimentali", level=1)

    # 4.1 Baseline
    doc.add_heading("4.1 Baseline — LLaMA 3.3:70b", level=2)
    doc.add_paragraph(
        "L'esperimento baseline utilizza il modello LLaMA 3.3:70b con il prompt pack default "
        "su tutti i 30 profili ufficiali. Risultati: 30/30 successo (100%), "
        "score aggregato 0.775 ± 0.049."
    )

    if BASELINE_LLAMA:
        model_key = "llama3.3:70b"
        rows = []
        for key, label in METRIC_NAMES:
            mean, std = get_metric(BASELINE_LLAMA, model_key, key)
            if mean is not None:
                std_str = f" ± {fmt(std)}" if std and std > 0 else ""
                rows.append((label, f"{fmt(mean)}{std_str}"))
        add_table(doc, ["Metrica", "Score (mean ± std)"], rows,
                  col_widths=[6, 6], bold_last_row=True)

    # 4.2 Ablation
    doc.add_heading("4.2 Ablation — Prompt Packs", level=2)
    doc.add_paragraph(
        "Lo studio di ablazione confronta l'effetto di tre varianti di system prompt "
        "(prompt pack) sulla qualità della generazione, mantenendo fisso il modello "
        "(LLaMA 3.3:70b) e il set di profili (n=30)."
    )

    rows = []
    for key, label in METRIC_NAMES:
        vals = []
        for summary, model_key in [
            (BASELINE_LLAMA, "llama3.3:70b"),
            (ABLATION_CONCISE, "llama3.3:70b"),
            (ABLATION_DETAILED, "llama3.3:70b"),
        ]:
            mean, _ = get_metric(summary, model_key, key)
            vals.append(fmt(mean) if mean is not None else "—")
        # Determine winner
        floats = []
        for v in vals:
            try:
                floats.append(float(v))
            except (ValueError, TypeError):
                floats.append(None)
        winner = ""
        valid = [(f, i) for i, f in enumerate(floats) if f is not None]
        if valid:
            best_idx = max(valid, key=lambda x: x[0])[1]
            winner = ["Default", "Concise", "Detailed"][best_idx]
        rows.append((label, vals[0], vals[1], vals[2], winner))

    add_table(doc,
              ["Metrica", "Default", "Concise", "Detailed", "Migliore"],
              rows, col_widths=[4.5, 2.5, 2.5, 2.5, 2.5], bold_last_row=True)

    doc.add_paragraph(
        "Il prompt pack default ottiene il miglior score aggregato (0.775 ± 0.049). "
        "Il pack concise sacrifica la coerenza narrativa (cross-scene coherence: 0.137 vs 0.588) "
        "pur migliorando il punteggio LLM Judge (0.967 vs 0.770). "
        "Il pack detailed aumenta la specificità (0.692 vs 0.535) ma introduce "
        "maggiore varianza (std 0.080 vs 0.049)."
    )

    # 4.3 Cross-Model
    doc.add_heading("4.3 Cross-Model — LLaMA 70B vs Qwen3 32B", level=2)
    doc.add_paragraph(
        "Il confronto cross-model valuta le differenze tra LLaMA 3.3:70b e Qwen3:32b "
        "sullo stesso set di profili con prompt pack default. Analisi statistica con "
        "test di Wilcoxon signed-rank (paired)."
    )

    rows = []
    llama_key = "llama3.3:70b"
    qwen_key = "qwen3:32b"
    for key, label in METRIC_NAMES:
        l_mean, l_std = get_metric(BASELINE_LLAMA, llama_key, key)
        q_mean, q_std = get_metric(BASELINE_QWEN, qwen_key, key)
        rows.append((label,
                      fmt(l_mean) if l_mean is not None else "—",
                      fmt(q_mean) if q_mean is not None else "—"))

    add_table(doc,
              ["Metrica", "LLaMA 3.3:70b", "Qwen3:32b"],
              rows, col_widths=[5, 4, 4], bold_last_row=True)

    doc.add_paragraph(
        "LLaMA 3.3:70b ottiene un aggregate score significativamente superiore "
        "(0.775 vs 0.729, p<.001, d=0.73). La differenza più marcata è nella "
        "cross-scene coherence (0.588 vs 0.102, d=5.46), dove Qwen mostra severe "
        "limitazioni nella generazione di narrative coese. Qwen eccelle nella "
        "specificità dei prompt (0.731 vs 0.535) e nell'auto-valutazione LLM Judge "
        "(0.996 vs 0.770)."
    )

    # 4.4 Per-Archetype
    doc.add_heading("4.4 Analisi per Archetipo", level=2)
    doc.add_paragraph(
        "L'analisi per archetipo (baseline LLaMA 3.3:70b, n=10 per archetipo) "
        "rivela differenze significative nelle performance tra i tre archetipi."
    )

    if BASELINE_LLAMA:
        rows = []
        for key, label in METRIC_NAMES:
            vals = []
            for arch in ["sage", "rebel", "lover"]:
                val = get_archetype_metric(BASELINE_LLAMA, arch, key)
                vals.append(fmt(val) if val is not None else "—")
            rows.append((label, vals[0], vals[1], vals[2]))

        add_table(doc,
                  ["Metrica", "Sage", "Rebel", "Lover"],
                  rows, col_widths=[5, 3, 3, 3], bold_last_row=True)

    doc.add_paragraph(
        "Sage ottiene il miglior score aggregato (0.805 ± 0.023), seguito da Lover "
        "(0.798 ± 0.018). Rebel presenta le performance più basse (0.723 ± 0.047) "
        "con significativa debolezza nell'archetype consistency (0.589) e nella "
        "specificità dei prompt (0.368). Questa asimmetria suggerisce che il modello "
        "fatica a generare contenuti dinamici/audaci mantenendo la coerenza tematica marina."
    )

    # 4.5 PolicyGate
    doc.add_heading("4.5 PolicyGate Compliance", level=2)
    doc.add_paragraph(
        "L'analisi della conformità PolicyGate (LLaMA 3.3:70b, default prompt) "
        "conferma la problematica dell'archetipo Rebel:"
    )

    policy_rows = [
        ("Green (PASS)", "8", "2", "9"),
        ("Yellow (WARN)", "2", "2", "0"),
        ("Red (FAIL)", "0", "6", "1"),
        ("Avg Violations", "0.0", "0.6", "0.1"),
    ]
    add_table(doc,
              ["Status", "Sage", "Rebel", "Lover"],
              policy_rows, col_widths=[4, 3, 3, 3])

    doc.add_paragraph(
        "Il 60% dei profili Rebel riceve flag RED, indicando violazioni delle policy "
        "del brand. Sage e Lover mostrano tassi di compliance significativamente "
        "superiori (80% e 90% GREEN rispettivamente)."
    )


def section_5_metriche(doc):
    doc.add_heading("5. Metriche di Valutazione", level=1)

    doc.add_heading("5.1 Metriche Automatiche (M_AUTO_01–13)", level=2)
    doc.add_paragraph(
        "Il framework comprende 13 metriche automatiche, ciascuna normalizzata "
        "nell'intervallo [0, 1], organizzate in quattro categorie:"
    )

    all_metrics = [
        ("M_AUTO_01", "Schema Compliance", "Strutturale",
         "Conformità alla struttura JSON richiesta (video_triptych, ost_prompt, metadata)"),
        ("M_AUTO_02", "Archetype Consistency", "Semantica",
         "Allineamento del contenuto al profilo archetipale dichiarato"),
        ("M_AUTO_03", "Role Sequence Valid", "Strutturale",
         "Correttezza della sequenza scene (start → evolve → end)"),
        ("M_AUTO_04", "Story Thread Presence", "Semantica",
         "Incorporazione dello story_thread_hint dal profilo utente"),
        ("M_AUTO_05", "Red Flag Score", "Qualitativa",
         "Assenza di termini nella blacklist PolicyGate (contenuti proibiti)"),
        ("M_AUTO_06", "Prompt Length Valid", "Strutturale",
         "Lunghezza dei prompt nel range ammesso (50-500 caratteri per scena)"),
        ("M_AUTO_07", "Archetype Lexical Fit", "Semantica",
         "Densità lessicale di vocabolario specifico per archetipo"),
        ("M_AUTO_08", "Cross-Scene Coherence", "Qualitativa",
         "Coerenza semantica tra scene consecutive del triptych"),
        ("M_AUTO_09", "Prompt Specificity", "Qualitativa",
         "Specificità e production-readiness dei prompt video"),
        ("M_AUTO_10", "Marine Vocabulary Ratio", "Semantica",
         "Rapporto vocabolario marino sul totale dei termini"),
        ("M_AUTO_11", "Narrative Coherence", "Qualitativa",
         "Consistenza del thread narrativo attraverso il triptych"),
        ("M_AUTO_12", "LLM Judge Quality", "LLM Judge",
         "Valutazione multi-dimensionale via LLM (5 dimensioni, scala 1-5, normalizzata)"),
        ("M_AUTO_13", "Pacing Progression", "Strutturale",
         "Progressione ritmica nell'arco narrativo (duration hints, intensità)"),
    ]
    add_table(doc,
              ["Codice", "Nome", "Categoria", "Descrizione"],
              [(m[0], m[1], m[2], m[3]) for m in all_metrics],
              col_widths=[2.5, 3.5, 2.5, 8.5])

    doc.add_heading("5.2 LLM-as-Judge", level=2)
    doc.add_paragraph(
        "La metrica M_AUTO_12 utilizza un modello LLM separato (qwen3:32b con /no_think "
        "per disabilitare il thinking mode, 3.4× più veloce) come giudice automatico. "
        "Il giudice valuta l'output su 5 dimensioni (scala 1-5):"
    )
    dims = [
        "Visual Clarity — produzione-readiness dei prompt video",
        "Archetype Alignment — coerenza con le caratteristiche dell'archetipo",
        "Narrative Coherence — fluidità narrativa attraverso il triptych",
        "Emotional Resonance — impatto emotivo del contenuto",
        "Marine Adherence — aderenza al tema marino/costiero",
    ]
    for d in dims:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("5.3 Framework per Valutazione Umana", level=2)
    doc.add_paragraph(
        "È stato predisposto un framework per la valutazione umana basato su 5 dimensioni "
        "con scala Likert (1-5), speculare alla valutazione automatica LLM Judge. "
        "Il framework è implementato nell'interfaccia Streamlit (pagina Evaluate) e "
        "consente la raccolta sistematica di giudizi umani per la validazione "
        "incrociata con le metriche automatiche."
    )
    doc.add_paragraph(
        "L'implementazione della campagna di valutazione umana è prevista come attività "
        "futura, previa definizione del protocollo sperimentale con il team NO NOISE."
    )


def section_6_sviluppi(doc):
    doc.add_heading("6. Indicazioni per Sviluppi Futuri", level=1)

    topics = [
        ("6.1 Profilazione Avanzata (NeedsProfiler LLM Agent)",
         "Evoluzione dal modello a 3 archetipi verso un sistema di profilazione dinamica "
         "basato su Big Five personality traits. Un agente NeedsProfiler costruirebbe "
         "profili utente attraverso interazioni conversazionali, consentendo una "
         "personalizzazione più granulare. Questo agente sostituirebbe la categorizzazione "
         "discreta con un embedding continuo nello spazio personalità."),

        ("6.2 Archivio Narrativo (Summarizer + Memory Store)",
         "Implementazione di un archivio narrativo per la continuità esperienziale: "
         "un Summarizer Agent comprimerebbe progressivamente le esperienze precedenti "
         "in summary compatti, mentre un Memory Store (vector database) consentirebbe "
         "il retrieval contestuale. Questo abilita la costruzione di \"storie\" "
         "che si sviluppano nel tempo con l'ospite."),

        ("6.3 Generazione Multi-Modale (Video + Musica + Diary)",
         "Estensione della pipeline di generazione per includere: generazione musicale "
         "da prompt OST, composizione di video + audio sincronizzati, e generazione "
         "di un \"diary\" narrativo testuale che accompagni l'esperienza visiva. "
         "L'obiettivo è una produzione multimediale completamente automatizzata."),

        ("6.4 Deployment On-Yacht (Edge Inference, Hybrid Storage)",
         "Architettura per il deployment reale su yacht: inferenza edge con modelli "
         "quantizzati per operatività offline, storage locale (SQLite) con "
         "sincronizzazione opportunistica verso cloud, gestione del budget "
         "computazionale in base alla connettività disponibile e alle risorse "
         "hardware a bordo."),
    ]
    for title, text in topics:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)


def appendix_a_profili(doc):
    doc.add_heading("Appendice A — Lista dei 30 Profili Ufficiali", level=1)

    profiles_dir = BASE_DIR / "data" / "profiles" / "official"
    rows = []
    for prefix, archetype in [("S", "Sage"), ("R", "Rebel"), ("L", "Lover")]:
        for i in range(1, 11):
            pid = f"{prefix}-{i:02d}"
            profile_path = profiles_dir / f"{pid}.json"
            genre, bpm, thread = "—", "—", "—"
            if profile_path.exists():
                try:
                    with open(profile_path) as f:
                        p = json.load(f)
                    up = p.get("user_profile", {})
                    ms = up.get("music_seed", {})
                    genre = ms.get("top_genre", "—")
                    bpm = str(ms.get("bpm", "—"))
                    thread = up.get("story_thread_hint", "—")
                except Exception:
                    pass
            rows.append((pid, archetype, genre, bpm, thread))

    add_table(doc,
              ["ID", "Archetipo", "Genere Musicale", "BPM", "Story Thread"],
              rows, col_widths=[2, 2.5, 3.5, 2, 5])


def appendix_b_metriche(doc):
    doc.add_heading("Appendice B — Dettaglio 13 Metriche Automatiche", level=1)

    metrics_detail = [
        ("M_AUTO_01 — Schema Compliance",
         "Valida la conformità dell'output alla struttura JSON richiesta. "
         "Verifica la presenza e il tipo corretto di: video_triptych (array di 3 oggetti con "
         "scene_role, prompt, duration_hint, mood_tags), ost_prompt (con prompt, genre, bpm), "
         "metadata (con archetype_detected, story_thread_used). Score binario per campo, "
         "media ponderata. Score: 0.0 (completamente non conforme) – 1.0 (perfettamente conforme)."),

        ("M_AUTO_02 — Archetype Consistency",
         "Valuta l'allineamento semantico del contenuto generato con le caratteristiche "
         "dell'archetipo dichiarato nel profilo. Analizza la densità di keyword specifiche "
         "per archetipo: Sage (contemplative, minimal, serene...), Rebel (dynamic, bold, "
         "powerful...), Lover (warm, intimate, sensual...). Score 0.0-1.0 basato sul "
         "rapporto keyword corrette / totali."),

        ("M_AUTO_03 — Role Sequence Valid",
         "Verifica la correttezza della sequenza dei ruoli nelle 3 scene del triptych. "
         "Sequenza attesa: [start, evolve, end]. Score binario: 1.0 se corretta, 0.0 altrimenti."),

        ("M_AUTO_04 — Story Thread Presence",
         "Misura l'incorporazione dello story_thread_hint dal profilo utente nel contenuto "
         "generato. Verifica la presenza di elementi tematici coerenti con l'hint narrativo "
         "nelle scene e nei metadata."),

        ("M_AUTO_05 — Red Flag Score",
         "Quantifica l'assenza di violazioni PolicyGate. Conta i termini nella blacklist "
         "(urban, violence, forest, ecc.) presenti nei prompt. Score 1.0 = nessuna violazione, "
         "decresce proporzionalmente al numero di termini proibiti trovati."),

        ("M_AUTO_06 — Prompt Length Valid",
         "Verifica che ogni prompt di scena rientri nel range ammesso (50-500 caratteri). "
         "Score per scena: 1.0 se nel range, 0.0 altrimenti. Media sulle 3 scene."),

        ("M_AUTO_07 — Archetype Lexical Fit",
         "Misura la densità lessicale di vocabolario specifico per l'archetipo. "
         "Rapporto tra token del vocabolario archetipale e token totali nei prompt video."),

        ("M_AUTO_08 — Cross-Scene Coherence",
         "Valuta la coerenza semantica tra scene consecutive usando similarità basata "
         "su embedding (TF-IDF + cosine similarity). Misura: sim(start,evolve) e "
         "sim(evolve,end), media delle due."),

        ("M_AUTO_09 — Prompt Specificity",
         "Misura la production-readiness dei prompt video. Penalizza: linguaggio vago, "
         "metafore, narrazione in prima persona ('we see', 'we witness'), riferimenti "
         "audio. Premia: dettagli concreti di camera, illuminazione, soggetto."),

        ("M_AUTO_10 — Marine Vocabulary Ratio",
         "Rapporto tra termini del vocabolario marino (sea, ocean, wave, shore, horizon, "
         "tide, yacht, sunset...) e il totale dei termini nei prompt. Garantisce il "
         "mantenimento del focus marino/costiero."),

        ("M_AUTO_11 — Narrative Coherence (SCORE)",
         "Score composito di coerenza narrativa che misura la consistenza del thread "
         "attraverso il triptych. Combina: ripetizione tematica, progressione emotiva, "
         "e coerenza delle mood_tags tra scene."),

        ("M_AUTO_12 — LLM Judge Quality",
         "Valutazione automatica tramite LLM giudice (qwen3:32b, /no_think). "
         "5 dimensioni (1-5): Visual Clarity, Archetype Alignment, Narrative Coherence, "
         "Emotional Resonance, Marine Adherence. Normalizzazione: (score-1)/4 → [0,1]."),

        ("M_AUTO_13 — Pacing Progression",
         "Misura la progressione del ritmo attraverso il triptych. Analizza i "
         "duration_hint e l'intensità percepita delle scene per verificare la presenza "
         "di un arco narrativo coerente con l'archetipo."),
    ]

    for title, desc in metrics_detail:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)


def appendix_c_esempio(doc):
    doc.add_heading("Appendice C — Esempio Output Director Completo", level=1)
    doc.add_paragraph(
        "Di seguito un esempio di output del Director per il profilo S-01 (Sage, ambient, "
        "60 BPM, story thread: single_cloud_gap) generato con LLaMA 3.3:70b:"
    )

    # Try to load an actual result
    results_path = BASE_DIR / "data" / "experiments" / "baseline_30_llama70b_v3" / "results.jsonl"
    example_output = None
    if results_path.exists():
        with open(results_path) as f:
            for line in f:
                try:
                    r = json.loads(line)
                    if r.get("profile_id") == "S-01" and r.get("success"):
                        example_output = r.get("output", {})
                        break
                except json.JSONDecodeError:
                    continue

    if example_output:
        formatted = json.dumps(example_output, indent=2, ensure_ascii=False)
        # Split into manageable chunks for Word
        p = doc.add_paragraph()
        p.style = "No Spacing"
        run = p.add_run(formatted[:3000])  # Truncate if very long
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        if len(formatted) > 3000:
            p = doc.add_paragraph()
            run = p.add_run("[... output troncato per brevità ...]")
            run.italic = True
    else:
        doc.add_paragraph("[Output di esempio non disponibile — eseguire l'esperimento baseline]")


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_style(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover(doc)
    section_1_sommario(doc)
    doc.add_page_break()
    section_2_attivita(doc)
    doc.add_page_break()
    section_3_architettura(doc)
    doc.add_page_break()
    section_4_risultati(doc)
    doc.add_page_break()
    section_5_metriche(doc)
    doc.add_page_break()
    section_6_sviluppi(doc)
    doc.add_page_break()
    appendix_a_profili(doc)
    doc.add_page_break()
    appendix_b_metriche(doc)
    doc.add_page_break()
    appendix_c_esempio(doc)

    out_path = OUT_DIR / "D3_Relazione_Finale_DRAFT.docx"
    doc.save(str(out_path))
    print(f"✓ D3 saved to: {out_path}")
    print(f"  Size: {out_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
